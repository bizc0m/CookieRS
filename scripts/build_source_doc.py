from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "CookieRS_Source_Guide.docx"


def add_title(document, title, subtitle):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(title)
    run.bold = True
    run.font.size = Pt(24)

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(subtitle)
    run.font.size = Pt(11)


def add_heading(document, text):
    document.add_heading(text, level=1)


def add_bullets(document, items):
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def add_code(document, text):
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)


def build():
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = document.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)

    add_title(document, "CookieRS - Guide source", "Application Electron open source inspiree de Biscuit/Cookie")

    add_heading(document, "Architecture")
    document.add_paragraph(
        "CookieRS est une application Electron compacte. Le process principal gere la fenetre native, "
        "les permissions Notification et les IPC. Le renderer contient l'interface, l'etat local, les webviews, "
        "les menus, les raccourcis, les skins et l'import/export JSON."
    )
    add_bullets(
        document,
        [
            "src/main.js : fenetre Electron, permissions et handlers IPC.",
            "src/preload.cjs : API exposee au renderer via contextBridge.",
            "src/renderer/app.js : etat, rendu UI, webviews, menus, partage, notifications.",
            "src/renderer/styles.css : layout, densites A/B/C, skins et panneau de reglages.",
            "docs/SOURCE.md : documentation technique Markdown.",
        ],
    )

    add_heading(document, "Fonctions principales")
    add_bullets(
        document,
        [
            "Workspaces groupes avec sessions separees par app et par groupe.",
            "Colonne app repliable en mode icones seules.",
            "Compteurs de notifications par app, via API Notification et fallback titre (N).",
            "Apps et groupes en drag and drop.",
            "Onglets secrets avec raccourcis Cmd/Ctrl+Shift+S et Cmd/Ctrl+Shift+H.",
            "Partage du texte selectionne avec URL vers copie, X, LinkedIn, Buffer ou Mail.",
            "Export et import complet de configuration JSON.",
        ],
    )

    add_heading(document, "Interface Cookie/Biscuit")
    document.add_paragraph(
        "Le lot 5 rapproche l'interface des captures fournies : sidebar plus simple, lignes app avec bouton "
        "options, footer + Groupe / OFF / Donate / + App, et panneau Parametres plein ecran."
    )
    add_bullets(
        document,
        [
            "La roue de reglages centralise les reglages globaux.",
            "Le deuxieme bouton de la fenetre ouvre les actions normales de page.",
            "Le masquage d'URL affiche seulement le domaine dans l'ecran principal.",
            "Les proprietes d'app et de groupe acceptent couleur, highlight, URL d'image et upload image.",
        ],
    )

    add_heading(document, "Etat persiste")
    document.add_paragraph("L'etat est stocke dans localStorage sous la cle cookiers.state.")
    add_code(
        document,
        "density, skin, customSkin, settingsOpen, settingsSection, maskUrl, sidebarCollapsed, "
        "secretsHidden, showHiddenApps, workspaces, appsByWorkspace, activeAppByWorkspace, tabsByApp",
    )

    add_heading(document, "Schema groupe")
    add_code(document, "id, name, icon, iconImage, color, highlightColor")

    add_heading(document, "Schema app")
    add_code(document, "id, name, url, color, highlightColor, iconImage, notifications, notificationCount, hidden")

    add_heading(document, "Raccourcis")
    add_bullets(
        document,
        [
            "Cmd/Ctrl+1..9 : changer de groupe.",
            "Alt+Cmd/Ctrl+Gauche/Droite : groupe precedent/suivant.",
            "Cmd/Ctrl+Shift+S : rendre l'onglet actif secret.",
            "Cmd/Ctrl+Shift+H : cacher ou afficher les onglets secrets.",
        ],
    )

    add_heading(document, "Build multi-OS")
    document.add_paragraph("CookieRS utilise electron-builder.")
    add_code(
        document,
        "npm install\nnpm start\nnpm run pack\nnpm run build:mac\nnpm run build:win\nnpm run build:linux",
    )
    document.add_paragraph(
        "La compilation finale doit idealement etre lancee sur chaque OS cible ou via CI multi-OS."
    )

    add_heading(document, "Limites connues")
    add_bullets(
        document,
        [
            "Les notifications dependent de la permission du site et de son usage reel de l'API Notification.",
            "Le bloqueur de publicites est reserve dans l'interface mais pas encore connecte a une liste de filtres.",
            "La synchro cloud est reservee dans les reglages mais reste locale pour l'instant.",
        ],
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
